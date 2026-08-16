import os
import glob
import re
import json
import struct
import sys
import olefile
from deep_translator import GoogleTranslator
from supabase import create_client, Client
from dotenv import load_dotenv

sys.stdout.reconfigure(encoding='utf-8')
load_dotenv('.env.local')

SUPABASE_URL = os.getenv('NEXT_PUBLIC_SUPABASE_URL')
SUPABASE_KEY = os.getenv('SUPABASE_SERVICE_ROLE_KEY')

if not SUPABASE_URL or not SUPABASE_KEY:
    print("❌ Missing Supabase credentials in .env.local")
    sys.exit(1)

supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

def extract_doc_raw(filepath):
    if filepath.endswith('.docx'):
        import docx
        doc = docx.Document(filepath)
        txt = '\n'.join([p.text for p in doc.paragraphs])
        for t in doc.tables:
            for r in t.rows:
                txt += '\n' + '\x07'.join([c.text.strip() for c in r.cells]) + '\x07\r'
        return txt
    try:
        ole = olefile.OleFileIO(filepath)
        word_stream = ole.openstream('WordDocument').read()
        flags = struct.unpack_from('<H', word_stream, 0x000A)[0]
        table_stream = ole.openstream('1Table' if (flags & 0x0200) else '0Table').read()
        fcClx = struct.unpack_from('<I', word_stream, 0x01A2)[0]
        lcbClx = struct.unpack_from('<I', word_stream, 0x01A6)[0]
        clx = table_stream[fcClx : fcClx + lcbClx]
        pos = 0
        while pos < len(clx):
            clxt = clx[pos]
            pos += 1
            if clxt == 1:
                pos += 2 + struct.unpack_from('<H', clx, pos)[0]
            elif clxt == 2:
                cb = struct.unpack_from('<I', clx, pos)[0]
                pos += 4
                plc = clx[pos : pos + cb]
                n = (cb - 4) // 12
                cps = [struct.unpack_from('<I', plc, i * 4)[0] for i in range(n + 1)]
                pcds_offset = (n + 1) * 4
                full = []
                for i in range(n):
                    fc_val = struct.unpack_from('<I', plc, pcds_offset + i * 8 + 2)[0]
                    cnt = cps[i+1] - cps[i]
                    fc = fc_val & 0x3FFFFFFF
                    if not (fc_val & 0x40000000):
                        full.append(word_stream[fc : fc + cnt * 2].decode('utf-16le', errors='ignore'))
                    else:
                        full.append(word_stream[fc//2 : fc//2 + cnt].decode('cp1252', errors='ignore'))
                ole.close()
                return ''.join(full)
        ole.close()
    except Exception as e:
        print(f"Error reading {filepath}: {e}")
    return ''

def clean(t):
    return re.sub(r'\s+', ' ', t).strip()

def parse_spec(filepath):
    raw = extract_doc_raw(filepath)
    filename = os.path.basename(filepath)
    data = {
        'file': filename,
        'filepath': filepath,
        'product_name': '',
        'packing': '',
        'usage': '',
        'storage_min': None,
        'storage_max': None,
        'shelf_life_months': None,
        'ingredients_text_en': '',
        'nutrition': {},
        'allergens': {}
    }
    
    # 1. Product Name
    m_name = re.search(r'Product name\s*:\s*([^|\n\r\x07]+)', raw, re.IGNORECASE)
    if m_name:
        data['product_name'] = clean(m_name.group(1))
    else:
        m_name2 = re.search(r'Trade Mark\s*:\s*FO FOOD PRODUCTS\s*([^\n\r|\x07]+)', raw, re.IGNORECASE)
        if m_name2:
            data['product_name'] = clean(m_name2.group(1))
        else:
            # Fallback filename
            data['product_name'] = os.path.splitext(filename)[0]

    # Clean product name
    data['product_name'] = re.sub(r'^\d+[\.-]\s*', '', data['product_name']).strip()
    
    # 2. Packing
    m_pack = re.search(r'Packing type\s*:\s*([^|\n\r\x07]+)', raw, re.IGNORECASE)
    if m_pack:
        data['packing'] = clean(m_pack.group(1))
        
    # 3. Usage
    m_usage = re.search(r'Type of Usage\s*:\s*([^|\n\r\x07]+)', raw, re.IGNORECASE)
    if m_usage:
        data['usage'] = clean(m_usage.group(1))
        
    # 4. Storage
    m_temp = re.search(r'(\d+)\s*[-–]\s*(\d+)\s*°?\s*C', raw)
    if m_temp:
        data['storage_min'] = int(m_temp.group(1))
        data['storage_max'] = int(m_temp.group(2))
        
    # 5. Shelf life
    m_shelf = re.search(r'shelf life of\s*(\d+)\s*years?', raw, re.IGNORECASE)
    if m_shelf:
        data['shelf_life_months'] = int(m_shelf.group(1)) * 12
    else:
        m_shelf2 = re.search(r'shelf life of\s*(\d+)\s*months?', raw, re.IGNORECASE)
        if m_shelf2:
            data['shelf_life_months'] = int(m_shelf2.group(1))
            
    # 6. Ingredients
    p1 = re.search(r'2\.\s*PRODUCT INGREDIENTS|PRODUCT INGREDIENTS|2\.\s*INGREDIENTS', raw, re.IGNORECASE)
    p2 = re.search(r'3\.\s*NUTRITIONAL|3\.\s*NUTRITION|NUTRITIONAL INFORMATION', raw, re.IGNORECASE)
    
    if p1 and p2 and p2.start() > p1.end():
        ing_section = raw[p1.end():p2.start()]
        cells = [c.strip() for c in re.split(r'[\x07\r\n]+', ing_section) if c.strip() and c.strip() not in ['—', '–', '-', '…', ':', '']]
        clean_cells = [c for c in cells if c.upper() not in ['INGREDIENTS', 'CLASS', 'EC CODE', 'MAX. DOSAGE', 'MAXIMUM DOSAGE', 'MAX.DOSAGE', 'PERCENTAGE', 'TOTAL', 'TOPLAM', 'QS', 'Q.S.']]
        
        # Build readable ingredients list
        # Items are usually: [Name, Class, EC, Dosage, Percentage] or [Name, Percentage]
        # Let's filter and join nicely
        items = []
        skip_words = {'FLAVORING', 'ACIDITY REGULATOR', 'COLORANT', 'PRESERVATIVE', 'THICKENER', 'EMULSIFIER', 'SWEETENER', 'HUMECTANT', 'BULKING AGENT', 'ANTIOXIDANT'}
        for c in clean_cells:
            # Check if this cell is purely a percentage or E-code or class
            if c.upper() in skip_words:
                continue
            if c.startswith('%') or c.endswith('%') or re.match(r'^\d+([,\.]\d+)?\s*%$', c):
                if items:
                    items[-1] += f" ({c.strip()})"
                else:
                    items.append(c)
            elif re.match(r'^E\s*\d+[a-z]?(\s*,\s*E\s*\d+[a-z]?)*$', c, re.IGNORECASE):
                if items:
                    items[-1] += f" [{c.strip()}]"
                else:
                    items.append(c)
            elif len(c) > 1 and not re.match(r'^\d+([,\.]\d+)?\s*(gr|g|mg)/?(lt|kg)?$', c, re.IGNORECASE):
                items.append(c)
                
        data['ingredients_text_en'] = ", ".join(items)
        
    # 7. Nutrition facts
    nut = {}
    m_energy = re.search(r'Energy\s*\|?\s*(\d+)\s*kj\s*/\s*(\d+)\s*kcal', raw, re.IGNORECASE)
    if m_energy:
        nut['energie_kj'] = float(m_energy.group(1))
        nut['energie_kcal'] = float(m_energy.group(2))
    else:
        m_kcal = re.search(r'(\d+)\s*kcal', raw, re.IGNORECASE)
        if m_kcal:
            nut['energie_kcal'] = float(m_kcal.group(1))
            
    m_fat = re.search(r'Fat\s*\|?\s*(\d+[.,]?\d*)\s*g', raw, re.IGNORECASE)
    if m_fat:
        nut['fett'] = float(m_fat.group(1).replace(',', '.'))
        
    m_sat = re.search(r'(?:saturates|-of which saturates)\s*\|?\s*(\d+[.,]?\d*)\s*g', raw, re.IGNORECASE)
    if m_sat:
        nut['davon_gesaettigt'] = float(m_sat.group(1).replace(',', '.'))
        
    m_carb = re.search(r'(?:Total Carbohydrate|Carbohydrate)\s*\|?\s*(\d+[.,]?\d*)\s*g', raw, re.IGNORECASE)
    if m_carb:
        nut['kohlenhydrate'] = float(m_carb.group(1).replace(',', '.'))
        
    m_sugar = re.search(r'Sugar\s*\|?\s*(\d+[.,]?\d*)\s*g', raw, re.IGNORECASE)
    if m_sugar:
        nut['davon_zucker'] = float(m_sugar.group(1).replace(',', '.'))
        
    m_fiber = re.search(r'(?:Dietary fibre|fibre|fiber)\s*\|?\s*(\d+[.,]?\d*)\s*g', raw, re.IGNORECASE)
    if m_fiber:
        nut['ballaststoffe'] = float(m_fiber.group(1).replace(',', '.'))
        
    m_prot = re.search(r'Protein\s*\|?\s*(\d+[.,]?\d*)\s*g', raw, re.IGNORECASE)
    if m_prot:
        nut['eiweiss'] = float(m_prot.group(1).replace(',', '.'))
        
    m_salt = re.search(r'(?:Sodium/salt|Salt|Sodium)\s*\|?\s*(?:(\d+[.,]?\d*)\s*g\s*/\s*)?(\d+[.,]?\d*)\s*g', raw, re.IGNORECASE)
    if m_salt:
        val = m_salt.group(2) or m_salt.group(1)
        if val:
            nut['salz'] = float(val.replace(',', '.'))
            
    if nut:
        data['nutrition'] = {'pro_100g': nut}
        
    # 8. Allergens
    m_all = re.search(r'ALLERGEN ALERT:(.*?)(?=9\.|GMO DECLARATION|9\. GMO|10\.)', raw, re.DOTALL | re.IGNORECASE)
    all_text = m_all.group(1).lower() if m_all else ''
    
    is_free = 'does not contain' in all_text or 'free' in all_text or 'no allergen' in all_text or 'not contain any allergen' in all_text
    has_milk = 'milk' in all_text or 'dairy' in all_text or 'laktose' in all_text or 'whey' in all_text
    has_gluten = 'gluten' in all_text or 'wheat' in all_text
    has_soy = 'soy' in all_text or 'soja' in all_text
    has_nuts = 'nut' in all_text or 'hazelnut' in all_text or 'pistachio' in all_text or 'almond' in all_text
    
    if has_milk or has_gluten or has_soy or has_nuts:
        is_free = False
    elif not is_free and not (has_milk or has_gluten or has_soy or has_nuts):
        is_free = True # default for fruit syrups without allergen warnings

    all_obj = {
        'allergen_free': is_free,
        'milch': has_milk,
        'gluten': has_gluten,
        'soja': has_soy,
        'nuesse': has_nuts
    }
    
    if is_free:
        all_obj['contains_de'] = 'Allergenfrei / Keine Allergene enthalten.'
        all_obj['contains_tr'] = 'Alerjen içermez.'
        all_obj['contains_en'] = 'Allergen free.'
        all_obj['contains_ar'] = 'خالٍ من مسببات الحساسية.'
    else:
        parts_de = []
        parts_tr = []
        parts_en = []
        parts_ar = []
        if has_milk:
            parts_de.append('Milchprodukte')
            parts_tr.append('Süt ürünleri')
            parts_en.append('Dairy / Milk products')
            parts_ar.append('منتجات الألبان')
        if has_soy:
            parts_de.append('Sojaprodukte')
            parts_tr.append('Soya ürünleri')
            parts_en.append('Soy products')
            parts_ar.append('منتجات الصويا')
        if has_gluten:
            parts_de.append('Glutenhaltiges Getreide')
            parts_tr.append('Gluten')
            parts_en.append('Gluten')
            parts_ar.append('الغلوتين')
        if has_nuts:
            parts_de.append('Schalenfrüchte / Nüsse')
            parts_tr.append('Fındık / Fıstık / Kuruyemiş')
            parts_en.append('Tree nuts')
            parts_ar.append('المكسرات')
            
        all_obj['contains_de'] = f"Enthält {', '.join(parts_de)}."
        all_obj['contains_tr'] = f"{', '.join(parts_tr)} içerir."
        all_obj['contains_en'] = f"Contains {', '.join(parts_en)}."
        all_obj['contains_ar'] = f"يحتوي على {', '.join(parts_ar)}."
        
    data['allergens'] = all_obj
    return data

def main():
    spec_files = glob.glob('dokuments/FO Ürün Spektleri/**/*.doc*', recursive=True)
    # also get .doc without asterisk if needed, but .doc* catches .doc and .docx
    # remove temporary ~ files
    spec_files = [f for f in spec_files if not os.path.basename(f).startswith('~')]
    
    print(f"Found {len(spec_files)} spec files.")
    
    parsed_specs = []
    for sf in spec_files:
        spec = parse_spec(sf)
        if spec['product_name']:
            parsed_specs.append(spec)
            
    print(f"Successfully parsed {len(parsed_specs)} specs.")
    
    with open('parsed_all_specs.json', 'w', encoding='utf-8') as f:
        json.dump(parsed_specs, f, ensure_ascii=False, indent=2)
        
    print("Saved to parsed_all_specs.json")

if __name__ == '__main__':
    main()
