import os
import re
import json
import docx
import sys
from deep_translator import GoogleTranslator

sys.stdout.reconfigure(encoding='utf-8')

def clean_text(text):
    return " ".join(text.split()).strip()

def extract_from_docx(filepath):
    try:
        doc = docx.Document(filepath)
    except Exception as e:
        print(f"Error reading {filepath}: {e}")
        return None

    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
        
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(clean_text(cell.text))
            full_text.append(" | ".join(row_data))

    # Convert to single string for some regex searches
    text_content = "\n".join(full_text)

    data = {
        "orjinal_ad": "",
        "ad": {"en": "", "de": "", "tr": "", "ar": ""},
        "aciklama": {"en": "", "de": "", "tr": "", "ar": ""},
        "inhaltsstoffe": {"en": "", "de": "", "tr": "", "ar": ""},
        "naehrwerte": {},
        "allergene": {},
        "lagertemperatur_min_celsius": None,
        "lagertemperatur_max_celsius": None,
        "haltbarkeit_monate": None,
    }

    # Extract Product Name
    name_match = re.search(r"Product name\s*:\s*(.*)", text_content, re.IGNORECASE)
    if name_match:
        data["orjinal_ad"] = clean_text(name_match.group(1))
        data["ad"]["en"] = data["orjinal_ad"]

    # Extract Usage (to be used as description)
    usage_match = re.search(r"Type of Usage\s*:\s*(.*)", text_content, re.IGNORECASE)
    if usage_match:
        data["aciklama"]["en"] = clean_text(usage_match.group(1))

    # Extract Storage Conditions
    storage_match = re.search(r"Storing Conditions\s*:\s*(.*)", text_content, re.IGNORECASE)
    if storage_match:
        storage_text = storage_match.group(1)
        temp_match = re.search(r"(\d+)\s*-\s*(\d+)\s*?\s*C", storage_text)
        if temp_match:
            data["lagertemperatur_min_celsius"] = int(temp_match.group(1))
            data["lagertemperatur_max_celsius"] = int(temp_match.group(2))

    # Extract Shelf Life
    shelf_match = re.search(r"shelf life of\s*(\d+)\s*years", text_content, re.IGNORECASE)
    if shelf_match:
        data["haltbarkeit_monate"] = int(shelf_match.group(1)) * 12

    # Extract Allergens
    allergen_match = re.search(r"ALLERGEN ALERT:(.*?)(?=9\.|GMO DECLARATION)", text_content, re.IGNORECASE | re.DOTALL)
    allergens_text = ""
    if allergen_match:
        allergens_text = allergen_match.group(1).lower()
        if "milk" in allergens_text:
            data["allergene"]["milch"] = True
        if "soy" in allergens_text:
            data["allergene"]["soja"] = True
        if "nut" in allergens_text or "hazelnut" in allergens_text:
            data["allergene"]["nuesse"] = True
        if "gluten" in allergens_text or "wheat" in allergens_text:
            data["allergene"]["gluten"] = True

    # Parse Tables for Ingredients and Nutritional Info
    ingredients = []
    nutritional = {
        "energie_kj": 0, "energie_kcal": 0, "fett": 0, "davon_gesaettigt": 0,
        "kohlenhydrate": 0, "davon_zucker": 0, "ballaststoffe": 0, "eiweiss": 0, "salz": 0
    }

    for table in doc.tables:
        # Check if table is ingredients table (headers usually have INGREDIENTS)
        is_ingredient_table = False
        is_nutritional_table = False

        if len(table.rows) > 0 and len(table.columns) > 0:
            header_text = table.rows[0].cells[0].text.upper()
            if "INGREDIENT" in header_text:
                is_ingredient_table = True
            elif "ENERGY" in header_text or "ENERGY" in table.rows[0].cells[0].text.upper():
                # Wait, nutritional is often not a table, or is a 2 column table
                is_nutritional_table = True

        if is_ingredient_table:
            for row in table.rows[1:]:
                if len(row.cells) > 0:
                    ing_name = clean_text(row.cells[0].text)
                    if ing_name and ing_name.lower() not in ['total', 'toplam']:
                        ingredients.append(ing_name)

        # Let's try to extract nutritional from text instead since tables can be weird
        pass

    if ingredients:
        data["inhaltsstoffe"]["en"] = ", ".join(ingredients)

    # Fallback for ingredients if table parsing failed
    if not data["inhaltsstoffe"]["en"]:
        ing_match = re.search(r"PRODUCT INGREDIENTS\s*:(.*?)(?=3\.NUTRITIONAL)", text_content, re.IGNORECASE | re.DOTALL)
        if ing_match:
            lines = ing_match.group(1).split('\n')
            ings = []
            for line in lines:
                cleaned = clean_text(line)
                if cleaned and "INGREDIENTS" not in cleaned and "CLASS" not in cleaned and not re.match(r'^[\d,%]+$', cleaned) and cleaned != '':
                    if len(cleaned) > 2 and cleaned.lower() not in ['qs', 'e 150a', 'e 334', 'e 202', 'e 211', '0,25 gr/ lt', '0,15 gr/lt']:
                        # very basic heuristic
                        if not re.search(r'^\d+$', cleaned):
                             ings.append(cleaned)
            # data["inhaltsstoffe"]["en"] = ", ".join(ings) # Might be messy, better rely on table

    # Nutritional info from text using regex
    nutri_map = {
        r"Energy.*?(\d+)\s*kj\s*/\s*(\d+)\s*kcal": ("energie_kj", "energie_kcal"),
        r"Fat.*?(\d+[.,]?\d*)\s*g": ("fett",),
        r"saturates.*?(\d+[.,]?\d*)\s*g": ("davon_gesaettigt",),
        r"Carbohydrate.*?(\d+[.,]?\d*)\s*g": ("kohlenhydrate",),
        r"Sugar.*?(\d+[.,]?\d*)\s*g": ("davon_zucker",),
        r"fibre.*?(\d+[.,]?\d*)\s*g": ("ballaststoffe",),
        r"Protein.*?(\d+[.,]?\d*)\s*g": ("eiweiss",),
        r"Sodium/salt.*?(\d+[.,]?\d*)\s*g\s*/\s*(\d+[.,]?\d*)\s*g": ("sodium", "salz")
    }

    for pattern, keys in nutri_map.items():
        match = re.search(pattern, text_content, re.IGNORECASE)
        if match:
            for i, key in enumerate(keys):
                val_str = match.group(i+1).replace(',', '.')
                try:
                    if key != "sodium":
                         nutritional[key] = float(val_str)
                except ValueError:
                    pass

    # Assign parsed nutritional
    data["naehrwerte"] = {"pro_100g": nutritional}

    return data

def translate_dict(text_dict, source_lang='en'):
    if not text_dict.get(source_lang):
        return text_dict
        
    text = text_dict[source_lang]
    try:
        text_dict['de'] = GoogleTranslator(source=source_lang, target='de').translate(text)
        text_dict['tr'] = GoogleTranslator(source=source_lang, target='tr').translate(text)
        text_dict['ar'] = GoogleTranslator(source=source_lang, target='ar').translate(text)
    except Exception as e:
        safe_text = text[:20].encode('ascii', 'replace').decode('ascii')
        print(f"Translation error for '{safe_text}...': {e}")
    return text_dict

def main():
    base_dir = r"dokuments\FO Ürün Spektleri"
    results = []
    
    # Collect all docx
    docx_files = []
    for root, dirs, files in os.walk(base_dir):
        for f in files:
            if f.endswith(".docx") and not f.startswith("~$"):
                docx_files.append(os.path.join(root, f))
                
    print(f"Found {len(docx_files)} docx files.")
    
    # Test with just a few first for safety, or all?
    # User said "100'e yakin Word dosyasi", 75 files. Let's process all.
    count = 0
    for file in docx_files:
        count += 1
        safe_name = os.path.basename(file).encode('ascii', 'replace').decode('ascii')
        print(f"Processing {count}/{len(docx_files)}: {safe_name}")
        data = extract_from_docx(file)
        if data and data.get("orjinal_ad"):
            # Translate fields
            data["ad"] = translate_dict(data["ad"])
            data["aciklama"] = translate_dict(data["aciklama"])
            data["inhaltsstoffe"] = translate_dict(data["inhaltsstoffe"])
            
            results.append(data)
            
    with open("extracted_specs.json", "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
        
    print("Done. Saved to extracted_specs.json")

if __name__ == "__main__":
    main()
